"""
Stage 1: Document Conversion — DOCX/PDF → Markdown + text + metadata.

Converts raw deal documents into standardized markdown format with
table preservation, heading detection, and content hashing.

Ported from PayGen pipeline.ingestion.document_converter → backend.abs.ingestion
Import rewrite: pipeline.config.pipeline_config → backend.abs.config.pipeline_config
"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional


@dataclass
class ConversionResult:
    """Result of document conversion."""
    full_md_path: Path
    full_txt_path: Path
    metadata_path: Path
    content_hash: str
    page_count: int
    paragraph_count: int
    table_count: int
    heading_count: int
    source_format: str
    errors: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "full_md_path": str(self.full_md_path),
            "full_txt_path": str(self.full_txt_path),
            "metadata_path": str(self.metadata_path),
            "content_hash": self.content_hash,
            "page_count": self.page_count,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "heading_count": self.heading_count,
            "source_format": self.source_format,
            "errors": self.errors,
        }


def convert_document(
    source_path: Path,
    output_dir: Path,
    deal_id: str = "",
) -> ConversionResult:
    """
    Convert a DOCX or PDF document to Markdown + plain text + metadata.

    Args:
        source_path: Path to source .docx or .pdf file
        output_dir: Directory for output files (full.md, full.txt, metadata.json)
        deal_id: Deal identifier for metadata

    Returns:
        ConversionResult with paths and statistics

    Raises:
        FileNotFoundError: if source_path does not exist
        ValueError: if file format is unsupported
    """
    source_path = Path(source_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not source_path.exists():
        raise FileNotFoundError(f"Source document not found: {source_path}")

    suffix = source_path.suffix.lower()
    if suffix == ".docx":
        md_content, stats = _convert_docx(source_path)
    elif suffix == ".pdf":
        # Check config for PDF support
        try:
            from backend.abs.config.pipeline_config import get_config
            cfg = get_config()
            if not cfg.pdf.enabled:
                raise ValueError(
                    "PDF processing is disabled in pipeline_config.yaml. "
                    "Set pdf.enabled: true to enable."
                )
        except ImportError:
            pass  # Config not available, allow PDF conversion
        md_content, stats = _convert_pdf(source_path)
    elif suffix == ".md":
        md_content = source_path.read_text(encoding="utf-8")
        stats = {
            "page_count": 0,
            "paragraph_count": md_content.count("\n\n") + 1,
            "table_count": md_content.count("| ---"),
            "heading_count": len(re.findall(r"^#{1,6}\s", md_content, re.MULTILINE)),
            "source_format": "markdown",
        }
    elif suffix == ".txt":
        md_content = source_path.read_text(encoding="utf-8")
        stats = {
            "page_count": 0,
            "paragraph_count": md_content.count("\n\n") + 1,
            "table_count": 0,
            "heading_count": 0,
            "source_format": "text",
        }
    else:
        raise ValueError(
            f"Unsupported file format: '{suffix}'. "
            f"Supported: .docx, .pdf, .md, .txt"
        )

    content_hash = hashlib.sha256(md_content.encode("utf-8")).hexdigest()

    # Write full.md
    md_path = output_dir / "full.md"
    md_path.write_text(md_content, encoding="utf-8")

    # Write full.txt (stripped markdown)
    plain_text = _strip_markdown(md_content)
    txt_path = output_dir / "full.txt"
    txt_path.write_text(plain_text, encoding="utf-8")

    # Write metadata.json
    metadata = {
        "deal_id": deal_id,
        "source_file": source_path.name,
        "source_format": stats.get("source_format", suffix.lstrip(".")),
        "content_hash": content_hash,
        "page_count": stats.get("page_count", 0),
        "paragraph_count": stats.get("paragraph_count", 0),
        "table_count": stats.get("table_count", 0),
        "heading_count": stats.get("heading_count", 0),
        "char_count": len(md_content),
        "converted_at": datetime.now(timezone.utc).isoformat(),
    }
    metadata_path = output_dir / "metadata.json"
    metadata_path.write_text(
        json.dumps(metadata, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    return ConversionResult(
        full_md_path=md_path,
        full_txt_path=txt_path,
        metadata_path=metadata_path,
        content_hash=content_hash,
        page_count=stats.get("page_count", 0),
        paragraph_count=stats.get("paragraph_count", 0),
        table_count=stats.get("table_count", 0),
        heading_count=stats.get("heading_count", 0),
        source_format=stats.get("source_format", suffix.lstrip(".")),
    )


# ── DOCX Conversion ──────────────────────────────────────────

def _convert_docx(source_path: Path) -> tuple[str, dict]:
    """Convert DOCX to markdown with table and heading preservation."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx required for DOCX conversion. "
            "Install: pip install python-docx"
        )

    doc = Document(str(source_path))
    lines: list[str] = []
    table_count = 0
    heading_count = 0
    paragraph_count = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue
            paragraph_count += 1
            style = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                lines.append("")
                continue

            if "Heading" in style:
                level = _heading_level(style)
                heading_count += 1
                lines.append(f"\n{'#' * level} {text}\n")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
                lines.append("")

        elif tag == "tbl":
            table = _find_table(doc, element)
            if table is not None:
                table_count += 1
                md_table = _table_to_markdown(table)
                lines.append(md_table)
                lines.append("")

    md_content = "\n".join(lines)
    stats = {
        "page_count": len(doc.sections),
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "heading_count": heading_count,
        "source_format": "docx",
    }
    return md_content, stats


def _find_paragraph(doc, element):
    """Find the Paragraph object for a given XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Find the Table object for a given XML element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _heading_level(style_name: str) -> int:
    """Extract heading level from style name (e.g., 'Heading 2' → 2)."""
    match = re.search(r"(\d+)", style_name)
    if match:
        return min(int(match.group(1)), 6)
    return 1


def _table_to_markdown(table) -> str:
    """Convert a docx table to markdown table format."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    md_lines: list[str] = []
    # Header row
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad or truncate to match header width
        while len(row) < len(rows[0]):
            row.append("")
        md_lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

    return "\n".join(md_lines)


# ── PDF Conversion ────────────────────────────────────────────

def _convert_pdf(source_path: Path) -> tuple[str, dict]:
    """
    Convert PDF to markdown using pdfplumber with:
    - Font-size-based heading inference
    - Table extraction with deduplication
    - Page boundary markers
    - Configurable behavior via pipeline_config.yaml
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber required for PDF conversion. "
            "Install: pip install pdfplumber"
        )

    # Load config options
    detect_tables = True
    infer_headings = True
    try:
        from backend.abs.config.pipeline_config import get_config
        cfg = get_config()
        detect_tables = cfg.pdf.detect_tables
        infer_headings = cfg.pdf.infer_headings
    except Exception:
        pass

    lines: list[str] = []
    page_count = 0
    table_count = 0
    heading_count = 0
    paragraph_count = 0

    with pdfplumber.open(str(source_path)) as pdf:
        page_count = len(pdf.pages)

        for page_idx, page in enumerate(pdf.pages):
            # ── Extract tables first (to exclude their bboxes from text) ──
            table_bboxes: list[tuple] = []
            if detect_tables:
                tables = page.find_tables() or []
                for table_obj in tables:
                    table_data = table_obj.extract()
                    if table_data and any(
                        cell for row in table_data for cell in row if cell
                    ):
                        table_count += 1
                        md_table = _list_table_to_markdown(table_data)
                        if md_table:
                            lines.append(md_table)
                            lines.append("")
                    table_bboxes.append(table_obj.bbox)

            # ── Extract text (excluding table regions) ──
            page_for_text = page
            if table_bboxes and detect_tables:
                for bbox in table_bboxes:
                    try:
                        page_for_text = page_for_text.outside_bbox(bbox)
                    except Exception:
                        pass

            # Extract with character-level data for heading inference
            if infer_headings:
                chars = page_for_text.chars or []
                text_blocks = _extract_text_with_headings(chars)
                for block_text, is_heading, level in text_blocks:
                    if is_heading:
                        heading_count += 1
                        lines.append(f"\n{'#' * level} {block_text}\n")
                    else:
                        paragraph_count += 1
                        lines.append(block_text)
                        lines.append("")
            else:
                text = page_for_text.extract_text() or ""
                if text.strip():
                    paragraph_count += text.count("\n\n") + 1
                    processed = _infer_headings_from_text(text)
                    heading_count += processed.count("\n# ")
                    lines.append(processed)
                    lines.append("")

    md_content = "\n".join(lines)

    # Clean up excessive whitespace
    md_content = re.sub(r"\n{4,}", "\n\n\n", md_content)

    stats = {
        "page_count": page_count,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "heading_count": heading_count,
        "source_format": "pdf",
    }
    return md_content, stats


def _extract_text_with_headings(
    chars: list[dict],
) -> list[tuple[str, bool, int]]:
    """
    Group characters into text blocks and infer headings from font size.

    Returns list of (text, is_heading, heading_level) tuples.
    """
    if not chars:
        return []

    # Group chars into lines by y-coordinate
    line_groups: list[list[dict]] = []
    current_line: list[dict] = []
    current_top = None

    for ch in sorted(chars, key=lambda c: (round(c.get("top", 0), 1), c.get("x0", 0))):
        top = round(ch.get("top", 0), 1)
        if current_top is None or abs(top - current_top) > 3:
            if current_line:
                line_groups.append(current_line)
            current_line = [ch]
            current_top = top
        else:
            current_line.append(ch)

    if current_line:
        line_groups.append(current_line)

    if not line_groups:
        return []

    # Calculate font size statistics
    all_sizes = []
    for line_chars in line_groups:
        sizes = [c.get("size", 12) for c in line_chars if c.get("text", "").strip()]
        if sizes:
            all_sizes.append(sum(sizes) / len(sizes))

    if not all_sizes:
        return []

    median_size = sorted(all_sizes)[len(all_sizes) // 2]

    # Build text blocks
    results: list[tuple[str, bool, int]] = []
    for line_chars in line_groups:
        text = "".join(c.get("text", "") for c in line_chars).strip()
        if not text:
            continue

        sizes = [c.get("size", 12) for c in line_chars if c.get("text", "").strip()]
        avg_size = sum(sizes) / len(sizes) if sizes else median_size
        is_bold = any(
            "bold" in str(c.get("fontname", "")).lower()
            for c in line_chars
            if c.get("text", "").strip()
        )

        if avg_size > median_size * 1.3:
            level = 1 if avg_size > median_size * 1.6 else 2
            results.append((text, True, level))
        elif is_bold and avg_size > median_size * 1.1:
            results.append((text, True, 3))
        elif is_bold and text.isupper() and len(text) < 100:
            results.append((text, True, 3))
        else:
            results.append((text, False, 0))

    return results


def _infer_headings_from_text(text: str) -> str:
    """
    Infer heading levels from common legal document patterns in raw text.
    """
    heading_patterns = [
        (r"^(ARTICLE\s+[IVXLCDM]+[.\s]*.*?)$", 1),
        (r"^(ARTICLE\s+\d+[.\s]*.*?)$", 1),
        (r"^(Section\s+\d+\.\d+[.\s]*.*?)$", 2),
        (r"^(SECTION\s+\d+\.\d+[.\s]*.*?)$", 2),
        (r"^(EXHIBIT\s+[A-Z][.\s]*.*?)$", 2),
        (r"^(SCHEDULE\s+[A-Z0-9][.\s]*.*?)$", 2),
        (r"^([A-Z][A-Z\s]{5,60})$", 3),
    ]

    output_lines: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        matched = False
        for pattern, level in heading_patterns:
            m = re.match(pattern, stripped)
            if m:
                output_lines.append(f"\n{'#' * level} {m.group(1).strip()}\n")
                matched = True
                break
        if not matched:
            output_lines.append(line)

    return "\n".join(output_lines)


def _list_table_to_markdown(table: list[list]) -> str:
    """Convert a list-of-lists table to markdown."""
    if not table or not table[0]:
        return ""

    cleaned: list[list[str]] = []
    for row in table:
        cleaned.append([str(cell or "").strip().replace("|", "\\|") for cell in row])

    col_count = max(len(row) for row in cleaned)

    md_lines: list[str] = []
    header = cleaned[0]
    while len(header) < col_count:
        header.append("")
    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("| " + " | ".join(["---"] * col_count) + " |")

    for row in cleaned[1:]:
        while len(row) < col_count:
            row.append("")
        md_lines.append("| " + " | ".join(row[:col_count]) + " |")

    return "\n".join(md_lines)


# ── Markdown Stripping ────────────────────────────────────────

def _strip_markdown(md_text: str) -> str:
    """Strip markdown formatting to produce plain text."""
    text = md_text
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.*?)_{1,3}", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"^[\s-]+$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
